from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_from_directory, send_file
from models import db, Post, Category, User, Comment, Favorite, Follow, ReadingHistory
from werkzeug.security import check_password_hash, generate_password_hash
from sqlalchemy import or_, and_
import re
import unicodedata
import os
from datetime import datetime
from pypinyin import pinyin, Style
from werkzeug.utils import secure_filename
from io import BytesIO
from docx import Document

client_bp = Blueprint('client', __name__)

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'static', 'uploads')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def make_slug(title):
    normalized = unicodedata.normalize('NFKD', title)
    ascii_text = normalized.encode('ascii', 'ignore').decode('ascii')
    
    if ascii_text.strip():
        slug = re.sub(r'[^\w\s-]', '', ascii_text.lower())
        slug = re.sub(r'[\s_-]+', '-', slug)
        slug = slug.strip('-')
    else:
        slug = None
    
    if not slug:
        pinyin_list = pinyin(title, style=Style.NORMAL)
        slug = '-'.join([p[0] for p in pinyin_list])
        slug = re.sub(r'[^\w\s-]', '', slug)
        slug = re.sub(r'[\s_-]+', '-', slug)
        slug = slug.strip('-')
    
    if not slug:
        slug = 'post'
    
    return slug

@client_bp.route('/')
def index():
    return render_template('client/index.html')

@client_bp.route('/post/<slug>')
def post_detail(slug):
    post = Post.query.filter_by(slug=slug, is_published=True).first_or_404()
    post.views += 1
    db.session.commit()
    
    categories = Category.query.all()
    comments = Comment.query.filter_by(post_id=post.id, parent_id=None).order_by(Comment.created_at.desc()).all()
    
    user_id = session.get('user_id')
    is_favorited = False
    is_following = False
    if user_id:
        is_favorited = Favorite.query.filter_by(user_id=user_id, post_id=post.id).first() is not None
        if post.author_id and post.author_id != user_id:
            is_following = Follow.query.filter_by(user_id=user_id, author_id=post.author_id).first() is not None
    
    return render_template('client/post.html', post=post, categories=categories, comments=comments, 
                           is_favorited=is_favorited, is_following=is_following)

@client_bp.route('/category/<slug>')
def category(slug):
    category = Category.query.filter_by(slug=slug).first_or_404()
    page = request.args.get('page', 1, type=int)
    
    posts = Post.query.filter_by(category_id=category.id, is_published=True)\
        .order_by(Post.created_at.desc())\
        .paginate(page=page, per_page=10, error_out=False)
    
    categories = Category.query.all()
    return render_template('client/category.html', category=category, posts=posts, categories=categories)

@client_bp.route('/type/<post_type>')
def post_type(post_type):
    page = request.args.get('page', 1, type=int)
    
    posts = Post.query.filter_by(post_type=post_type, is_published=True)\
        .order_by(Post.created_at.desc())\
        .paginate(page=page, per_page=10, error_out=False)
    
    categories = Category.query.all()
    type_names = {'note': '教学笔记', 'resource': '教育资源', 'experience': '教育心得'}
    return render_template('client/type.html', posts=posts, categories=categories, 
                           post_type=post_type, type_name=type_names.get(post_type, post_type))

@client_bp.route('/search')
def search():
    keyword = request.args.get('keyword', '')
    page = request.args.get('page', 1, type=int)
    
    if keyword:
        posts = Post.query.filter(
            or_(Post.title.contains(keyword), Post.content.contains(keyword)),
            Post.is_published==True
        ).order_by(Post.created_at.desc()).paginate(page=page, per_page=10, error_out=False)
    else:
        posts = []
    
    categories = Category.query.all()
    return render_template('client/search.html', posts=posts, keyword=keyword, categories=categories)

@client_bp.route('/comment/<int:post_id>', methods=['POST'])
def add_comment(post_id):
    content = request.form.get('content')
    parent_id = request.form.get('parent_id', type=int)
    
    if content and session.get('user_id'):
        comment = Comment(
            content=content,
            post_id=post_id,
            user_id=session['user_id'],
            parent_id=parent_id
        )
        db.session.add(comment)
        db.session.commit()
        flash('评论成功！', 'success')
    else:
        flash('请先登录后再评论', 'warning')
    
    post = Post.query.get_or_404(post_id)
    return redirect(url_for('client.post_detail', slug=post.slug))

@client_bp.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['is_admin'] = user.is_admin
            flash('登录成功！', 'success')
            next_page = request.args.get('next')
            return redirect(next_page or url_for('client.index'))
        else:
            flash('用户名或密码错误', 'danger')
    
    return render_template('client/login.html')

@client_bp.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('username', None)
    session.pop('is_admin', None)
    flash('已退出登录', 'info')
    return redirect(url_for('client.index'))

@client_bp.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if not username or not email or not password:
            flash('所有字段都为必填项', 'danger')
            return redirect(url_for('client.register'))
        
        if password != confirm_password:
            flash('两次输入的密码不一致', 'danger')
            return redirect(url_for('client.register'))
        
        if len(password) < 6:
            flash('密码长度至少为6位', 'danger')
            return redirect(url_for('client.register'))
        
        existing_user = User.query.filter(
            (User.username == username) | (User.email == email)
        ).first()
        
        if existing_user:
            if existing_user.username == username:
                flash('用户名已存在', 'danger')
            else:
                flash('邮箱已被注册', 'danger')
            return redirect(url_for('client.register'))
        
        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
        
        new_user = User(
            username=username,
            email=email,
            password_hash=hashed_password
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        flash('注册成功！请登录', 'success')
        return redirect(url_for('client.login'))
    
    return render_template('client/register.html')

@client_bp.route('/about')
def about():
    categories = Category.query.all()
    return render_template('client/about.html', categories=categories)

@client_bp.route('/agent-center')
def agent_center():
    categories = Category.query.all()
    return render_template('client/agent_center.html', categories=categories)

@client_bp.route('/lesson-planner')
def lesson_planner():
    categories = Category.query.all()
    return render_template('client/lesson_planner.html', categories=categories)

@client_bp.route('/export/word', methods=['POST'])
def export_word():
    """导出教案或习题为Word文档"""
    try:
        data = request.get_json()
        content_type = data.get('content_type')  # 'lesson_plan' or 'quiz'
        content = data.get('content', '')
        title = data.get('title', '')
        
        if not content:
            return jsonify({'success': False, 'message': '没有内容可导出'}), 400
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        if content_type == 'lesson_plan':
            doc.add_heading('教案大纲', 0)
            if title:
                doc.add_heading(title, 1)
        else:
            doc.add_heading('练习题', 0)
            if title:
                doc.add_heading(title, 1)
        
        # 解析并添加内容
        lines = content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                current_paragraph = None
                continue
            
            # 检查是否为标题（以 # 开头或包含特定模式）
            if line.startswith('## '):
                doc.add_heading(line[3:], 2)
                current_paragraph = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
                current_paragraph = None
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
                current_paragraph = None
            # 检查是否为列表项
            elif line.startswith('- ') or line.startswith('* '):
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.style = 'List Bullet'
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or re.match(r'^\d+\. ', line):
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.style = 'List Number'
                match = re.match(r'^\d+\. (.+)', line)
                if match:
                    p = doc.add_paragraph(match.group(1), style='List Number')
            else:
                # 普通段落
                doc.add_paragraph(line)
                current_paragraph = None
        
        # 保存文档到内存缓冲区
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if content_type == 'lesson_plan':
            filename = f"教案_{timestamp}.docx"
        else:
            filename = f"练习题_{timestamp}.docx"
        
        # 返回下载响应
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'导出失败: {str(e)}'}), 500

@client_bp.route('/courseware-optimizer')
def courseware_optimizer():
    categories = Category.query.all()
    return render_template('client/courseware_optimizer.html', categories=categories)

@client_bp.route('/publish', methods=['GET', 'POST'])
def publish():
    if not session.get('user_id'):
        flash('请先登录后再发布文章', 'warning')
        return redirect(url_for('client.login', next=url_for('client.publish')))
    
    categories = Category.query.all()
    
    if request.method == 'POST':
        print(f"[DEBUG] 开始处理发布请求")
        title = request.form.get('title')
        content = request.form.get('content')
        summary = request.form.get('summary')
        post_type = request.form.get('post_type')
        category_id = request.form.get('category_id')
        is_published = 'is_published' in request.form
        cover_image = request.form.get('cover_image')
        
        print(f"[DEBUG] 表单数据: title={title}, post_type={post_type}, category_id={category_id}, is_published={is_published}")
        
        if not title or not content or not category_id:
            flash('标题、内容和分类为必填项', 'danger')
            return redirect(url_for('client.publish'))
        
        slug = make_slug(title)
        print(f"[DEBUG] 生成slug: {slug}")
        
        existing_slug = Post.query.filter_by(slug=slug).first()
        if existing_slug:
            slug = f"{slug}-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        
        new_post = Post(
            title=title,
            slug=slug,
            content=content,
            summary=summary,
            post_type=post_type,
            category_id=category_id,
            author_id=session['user_id'],
            is_published=is_published,
            cover_image=cover_image,
            views=0
        )
        
        db.session.add(new_post)
        db.session.commit()
        print(f"[DEBUG] 文章已保存到数据库, ID={new_post.id}, is_published={is_published}")
        
        # 验证文章是否真的被保存到数据库
        print(f"[DEBUG] 开始验证数据库提交...")
        try:
            # 重新查询刚刚创建的文章
            committed_post = Post.query.get(new_post.id)
            if committed_post:
                print(f"[DEBUG] 验证成功: 找到文章 ID={committed_post.id}, title={committed_post.title}")
                print(f"[DEBUG] 文章状态: is_published={committed_post.is_published}, slug={committed_post.slug}")
                
                # 检查首页会显示的文章数量
                published_posts = Post.query.filter_by(is_published=True).count()
                draft_posts = Post.query.filter_by(is_published=False).count()
                print(f"[DEBUG] 数据库统计: 已发布={published_posts}, 草稿={draft_posts}")
                
                if is_published:
                    flash('文章发布成功！', 'success')
                    return redirect(url_for('client.post_detail', slug=new_post.slug))
                else:
                    flash('文章已保存为草稿', 'success')
                    return redirect(url_for('client.index'))
            else:
                print(f"[DEBUG] 验证失败: 未找到文章 ID={new_post.id}")
                flash('文章保存失败，请重试', 'danger')
                return redirect(url_for('client.publish'))
        except Exception as e:
            print(f"[DEBUG] 验证过程出错: {str(e)}")
            flash(f'文章保存失败: {str(e)}', 'danger')
            return redirect(url_for('client.publish'))
    
    return render_template('client/publish.html', categories=categories)

@client_bp.route('/upload/image', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    
    file = request.files['image']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}")
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        return jsonify({
            'url': url_for('static', filename=f'uploads/{filename}'),
            'filename': filename
        })
    
    return jsonify({'error': '不支持的文件格式'}), 400

@client_bp.route('/uploads/<path:filename>')
def serve_upload(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

@client_bp.route('/api/favorites', methods=['GET'])
def get_favorites():
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    favorites = Favorite.query.filter_by(user_id=session['user_id'])\
        .order_by(Favorite.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'favorites': [{
            'id': f.id,
            'post_id': f.post_id,
            'title': f.post.title,
            'slug': f.post.slug,
            'summary': f.post.summary,
            'cover_image': f.post.cover_image,
            'created_at': f.created_at.strftime('%Y-%m-%d %H:%M')
        } for f in favorites.items],
        'total': favorites.total,
        'pages': favorites.pages
    })

@client_bp.route('/api/favorites/<int:post_id>', methods=['POST'])
def toggle_favorite(post_id):
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    existing = Favorite.query.filter_by(
        user_id=session['user_id'],
        post_id=post_id
    ).first()
    
    if existing:
        db.session.delete(existing)
        db.session.commit()
        return jsonify({'status': 'removed', 'message': '已取消收藏'})
    else:
        favorite = Favorite(
            user_id=session['user_id'],
            post_id=post_id
        )
        db.session.add(favorite)
        db.session.commit()
        return jsonify({'status': 'added', 'message': '收藏成功'})

@client_bp.route('/api/favorites/check/<int:post_id>', methods=['GET'])
def check_favorite(post_id):
    if not session.get('user_id'):
        return jsonify({'is_favorited': False})
    
    existing = Favorite.query.filter_by(
        user_id=session['user_id'],
        post_id=post_id
    ).first()
    
    return jsonify({'is_favorited': bool(existing)})

@client_bp.route('/favorites')
def favorites_page():
    if not session.get('user_id'):
        flash('请先登录后再查看收藏夹', 'warning')
        return redirect(url_for('client.login'))
    
    categories = Category.query.all()
    return render_template('client/favorites.html', categories=categories)

@client_bp.route('/api/follows', methods=['GET'])
def get_follows():
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    follows = Follow.query.filter_by(user_id=session['user_id']).all()
    
    following = []
    for follow in follows:
        if follow.author:
            following.append({
                'id': follow.id,
                'author_id': follow.author_id,
                'username': follow.author.username,
                'bio': follow.author.bio,
                'avatar': follow.author.avatar,
                'post_count': follow.author.posts.count()
            })
    
    all_authors = User.query.filter(User.id != session['user_id']).all()
    following_ids = [f.author_id for f in follows]
    available_authors = [a for a in all_authors if a.id not in following_ids]
    
    return jsonify({'following': following, 'available_authors': [{
        'id': a.id,
        'username': a.username,
        'bio': a.bio,
        'avatar': a.avatar,
        'post_count': a.posts.count()
    } for a in available_authors]})

@client_bp.route('/api/follows', methods=['POST'])
def toggle_follow():
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    data = request.get_json()
    author_id = data.get('author_id')
    
    if not author_id:
        return jsonify({'error': '缺少作者ID'}), 400
    
    if int(author_id) == session['user_id']:
        return jsonify({'error': '不能关注自己'}), 400
    
    existing = Follow.query.filter_by(
        user_id=session['user_id'],
        author_id=author_id
    ).first()
    
    if existing:
        db.session.delete(existing)
        db.session.commit()
        return jsonify({'status': 'unfollowed', 'message': '已取消关注'})
    else:
        follow = Follow(
            user_id=session['user_id'],
            author_id=author_id
        )
        db.session.add(follow)
        db.session.commit()
        return jsonify({'status': 'followed', 'message': '关注成功'})

@client_bp.route('/api/follows/check/<int:author_id>', methods=['GET'])
def check_follow(author_id):
    if not session.get('user_id'):
        return jsonify({'is_following': False})
    
    existing = Follow.query.filter_by(
        user_id=session['user_id'],
        author_id=author_id
    ).first()
    
    return jsonify({'is_following': bool(existing)})

@client_bp.route('/following')
def following_page():
    if not session.get('user_id'):
        flash('请先登录后再查看关注列表', 'warning')
        return redirect(url_for('client.login'))
    
    categories = Category.query.all()
    return render_template('client/following.html', categories=categories)

@client_bp.route('/api/reading-history', methods=['GET'])
def get_reading_history():
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    history = ReadingHistory.query.filter_by(user_id=session['user_id'])\
        .order_by(ReadingHistory.read_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'history': [{
            'id': h.id,
            'post_id': h.post_id,
            'title': h.post.title,
            'slug': h.post.slug,
            'summary': h.post.summary,
            'cover_image': h.post.cover_image,
            'read_at': h.read_at.strftime('%Y-%m-%d %H:%M'),
            'read_duration': h.read_duration
        } for h in history.items],
        'total': history.total,
        'pages': history.pages
    })

@client_bp.route('/api/reading-history', methods=['POST'])
def add_reading_history():
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    data = request.get_json()
    post_id = data.get('post_id')
    read_duration = data.get('read_duration', 0)
    
    if not post_id:
        return jsonify({'error': '缺少文章ID'}), 400
    
    existing = ReadingHistory.query.filter_by(
        user_id=session['user_id'],
        post_id=post_id
    ).first()
    
    if existing:
        existing.read_at = datetime.utcnow()
        existing.read_duration = read_duration
    else:
        history = ReadingHistory(
            user_id=session['user_id'],
            post_id=post_id,
            read_duration=read_duration
        )
        db.session.add(history)
    
    db.session.commit()
    return jsonify({'status': 'success'})

@client_bp.route('/api/reading-history/<int:history_id>', methods=['DELETE'])
def delete_reading_history(history_id):
    if not session.get('user_id'):
        return jsonify({'error': '请先登录'}), 401
    
    history = ReadingHistory.query.filter_by(
        id=history_id,
        user_id=session['user_id']
    ).first_or_404()
    
    db.session.delete(history)
    db.session.commit()
    return jsonify({'status': 'success'})

@client_bp.route('/reading-history')
def reading_history_page():
    if not session.get('user_id'):
        flash('请先登录后再查看阅读历史', 'warning')
        return redirect(url_for('client.login'))
    
    categories = Category.query.all()
    return render_template('client/reading_history.html', categories=categories)

@client_bp.route('/api/recommendations', methods=['GET'])
def get_recommendations():
    if not session.get('user_id'):
        posts = Post.query.filter_by(is_published=True)\
            .order_by(Post.views.desc())\
            .limit(6).all()
        return jsonify({
            'recommendations': [{
                'id': p.id,
                'title': p.title,
                'slug': p.slug,
                'summary': p.summary,
                'cover_image': p.cover_image,
                'views': p.views,
                'category_name': p.category.name
            } for p in posts]
        })
    
    user_id = session['user_id']
    
    followed_authors = db.session.query(Follow.author_id)\
        .filter(Follow.user_id == user_id)\
        .all()
    followed_author_ids = [a[0] for a in followed_authors]
    
    favorite_posts = db.session.query(Favorite.post_id)\
        .filter(Favorite.user_id == user_id).all()
    favorite_post_ids = [p[0] for p in favorite_posts]
    
    history_posts = db.session.query(ReadingHistory.post_id)\
        .filter(ReadingHistory.user_id == user_id).all()
    history_post_ids = [p[0] for p in history_posts]
    
    read_category_ids = db.session.query(Post.category_id)\
        .filter(Post.id.in_(history_post_ids)).all()
    read_category_ids = [c[0] for c in read_category_ids]
    
    followed_posts = Post.query.filter(
        Post.is_published == True,
        Post.author_id.in_(followed_author_ids)
    ).order_by(Post.created_at.desc()).limit(6).all() if followed_author_ids else []
    
    if read_category_ids:
        related_posts = Post.query.filter(
            Post.is_published == True,
            Post.category_id.in_(read_category_ids)
        ).order_by(Post.views.desc()).limit(6).all()
    else:
        related_posts = Post.query.filter(
            Post.is_published == True
        ).order_by(Post.views.desc()).limit(6).all()
    
    exclude_ids = set(favorite_post_ids + history_post_ids)
    recommended_posts = [p for p in related_posts if p.id not in exclude_ids]
    
    if len(recommended_posts) < 6:
        additional = Post.query.filter(
            Post.is_published == True,
            Post.id.not_in([p.id for p in recommended_posts] + list(exclude_ids))
        ).order_by(Post.created_at.desc()).limit(6 - len(recommended_posts)).all()
        recommended_posts.extend(additional)
    
    return jsonify({
        'subscribed': [{
            'id': p.id,
            'title': p.title,
            'slug': p.slug,
            'summary': p.summary,
            'cover_image': p.cover_image,
            'views': p.views,
            'category_name': p.category.name
        } for p in followed_posts[:6]],
        'recommended': [{
            'id': p.id,
            'title': p.title,
            'slug': p.slug,
            'summary': p.summary,
            'cover_image': p.cover_image,
            'views': p.views,
            'category_name': p.category.name
        } for p in recommended_posts[:6]]
    })

@client_bp.route('/international-news')
def international_news():
    from services.news_service import get_latest_news, get_topics, get_featured_news, get_news_by_source
    from models import InternationalEducation

    page = request.args.get('page', 1, type=int)
    selected_topic = request.args.get('topic')
    selected_source = request.args.get('source')

    if selected_source:
        news_items = InternationalEducation.query\
            .filter_by(source_name=selected_source, is_featured=True)\
            .order_by(InternationalEducation.publish_date.desc())\
            .paginate(page=page, per_page=10, error_out=False)
    elif selected_topic:
        news_items = InternationalEducation.query\
            .filter_by(topic=selected_topic, is_featured=True)\
            .order_by(InternationalEducation.publish_date.desc())\
            .paginate(page=page, per_page=10, error_out=False)
    else:
        news_items = InternationalEducation.query\
            .filter_by(is_featured=True)\
            .order_by(InternationalEducation.publish_date.desc())\
            .paginate(page=page, per_page=10, error_out=False)

    topics = get_topics()
    featured_news = get_featured_news(limit=5)

    sources = db.session.query(
        InternationalEducation.source_name,
        db.func.count(InternationalEducation.id)
    ).filter(InternationalEducation.is_featured == True)\
     .group_by(InternationalEducation.source_name)\
     .all()

    return render_template(
        'client/international_news.html',
        news_items=news_items.items,
        page=news_items.page,
        total_pages=news_items.pages,
        total_news=news_items.total,
        topics=topics,
        sources=[{'source_name': s[0]} for s in sources],
        featured_news=featured_news,
        selected_topic=selected_topic,
        selected_source=selected_source,
        current_topic=selected_topic,
        current_source=selected_source,
        total_topics=len(topics) if topics else 0,
        total_sources=len(sources) if sources else 0,
        featured_topics=[{'name': t['name'], 'slug': t['name'].lower().replace(' ', '-'), 'icon': '📚', 'count': InternationalEducation.query.filter_by(topic=t['name'], is_featured=True).count()} for t in (topics[:5] if topics else [])]
    )

@client_bp.route('/api/international-news/refresh', methods=['POST'])
def refresh_international_news():
    from services.news_service import crawl_and_update_news
    try:
        count = crawl_and_update_news()
        return jsonify({'status': 'success', 'count': count})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@client_bp.route('/api/international-news')
def get_international_news():
    from services.news_service import get_latest_news, get_topics
    from models import InternationalEducation

    page = request.args.get('page', 1, type=int)
    topic = request.args.get('topic')

    if topic:
        news_items = InternationalEducation.query\
            .filter_by(topic=topic, is_featured=True)\
            .order_by(InternationalEducation.publish_date.desc())\
            .paginate(page=page, per_page=10, error_out=False)
    else:
        news_items = InternationalEducation.query\
            .filter_by(is_featured=True)\
            .order_by(InternationalEducation.publish_date.desc())\
            .paginate(page=page, per_page=10, error_out=False)

    topics = get_topics()

    return jsonify({
        'news': [{
            'id': n.id,
            'title': n.title,
            'source_name': n.source_name,
            'source_url': n.source_url,
            'country': n.country,
            'topic': n.topic,
            'content': n.content,
            'publish_date': n.publish_date.strftime('%Y-%m-%d') if n.publish_date else None,
            'fetched_at': n.fetched_at.strftime('%Y-%m-%d %H:%M')
        } for n in news_items.items],
        'topics': topics,
        'total': news_items.total,
        'pages': news_items.pages,
        'current_page': news_items.page
    })

@client_bp.route('/learning-progress')
def learning_progress():
    """学习进度管理页面"""
    from datetime import datetime, timedelta
    import random
    
    categories = Category.query.all()
    
    # 生成示例学习进度数据
    has_progress = True
    overall_completion = random.randint(35, 75)
    completed_count = random.randint(8, 15)
    total_count = random.randint(18, 25)
    difficult_count = random.randint(3, 6)
    difficult_percentage = round((difficult_count / total_count) * 100)
    study_hours = round(random.uniform(20, 80), 1)
    weekly_hours = round(random.uniform(8, 15), 1)
    daily_avg_hours = round(study_hours / 7, 1)
    
    # 学习路径数据
    learning_path = []
    stages = [
        ('基础知识', random.randint(60, 100)),
        ('核心概念', random.randint(40, 90)),
        ('进阶应用', random.randint(30, 70)),
        ('实战项目', random.randint(10, 50)),
        ('综合提升', random.randint(5, 40))
    ]
    
    for i, (name, progress) in enumerate(stages):
        is_completed = progress >= 85
        is_difficult = 40 <= progress < 65
        
        learning_path.append({
            'name': name,
            'progress': progress,
            'completed': is_completed,
            'is_difficult': is_difficult,
            'knowledge_points_count': random.randint(5, 12),
            'last_study_time': datetime.now() - timedelta(days=random.randint(0, 7))
        })
    
    # 困难内容数据
    difficult_contents = []
    topics = ['函数式编程', '异步处理', '设计模式', '算法优化', '数据库查询', 'API集成']
    for i in range(min(difficult_count, len(topics))):
        difficult_contents.append({
            'id': f'difficult_{i+1}',
            'title': topics[i],
            'comprehension_rate': random.randint(25, 55),
            'error_count': random.randint(2, 8),
            'difficulty_level': random.choice(['high', 'medium', 'low'])
        })
    
    # 每日学习趋势数据（最近14天）
    days = [(datetime.now() - timedelta(days=i)).strftime('%m/%d') for i in range(13, -1, -1)]
    daily_data = [round(random.uniform(1.5, 4.5), 1) for _ in range(14)]
    daily_study_data = {
        'labels': days,
        'data': daily_data
    }
    
    # 知识点掌握分布
    mastery_distribution = [
        max(0, min(100, overall_completion + random.randint(-5, 5))),
        max(0, min(100, 100 - overall_completion + random.randint(-10, 10))),
        max(0, min(100, random.randint(5, 15)))
    ]
    
    # 改进建议
    suggestions = [
        {
            'content': '建议加强基础知识的复习，特别是函数式编程和异步处理部分',
            'sub_suggestions': [
                '每天安排30分钟专项练习',
                '完成相关练习题至少20道',
                '观看教学视频并做笔记'
            ]
        },
        {
            'content': '增加实战项目的练习时间',
            'sub_suggestions': [
                '每周完成一个小型项目',
                '参与开源项目贡献代码',
                '编写技术博客总结经验'
            ]
        },
        {
            'content': '优化学习方法，提高学习效率',
            'sub_suggestions': [
                '使用番茄工作法管理时间',
                '定期回顾已学内容',
                '建立知识思维导图'
            ]
        },
        {
            'content': '加强算法和数据结构的学习',
            'sub_suggestions': [
                '每日刷LeetCode题目',
                '重点练习动态规划和图论',
                '参加算法竞赛提升能力'
            ]
        }
    ]

    # 课程目录/学习进程数据
    course_units = [
        {
            'name': '第一单元',
            'lessons': [
                {'number': '1.1', 'title': '算法引论', 'status': 'completed'},
                {'number': '1.2', 'title': '算法复杂度分析', 'status': 'completed'},
                {'number': '1.3', 'title': '递归与分治思想', 'status': 'completed'}
            ]
        },
        {
            'name': '第二单元',
            'lessons': [
                {'number': '2.1', 'title': '递归法', 'status': 'completed'},
                {'number': '2.2', 'title': '分治法', 'status': 'completed'},
                {'number': '2.3', 'title': '动态规划基础', 'status': 'completed'}
            ]
        },
        {
            'name': '第三单元',
            'lessons': [
                {'number': '3.1', 'title': '回溯法', 'status': 'completed'},
                {'number': '3.2', 'title': '分支限界法', 'status': 'in-progress'},
                {'number': '3.3', 'title': '剪枝优化技巧', 'status': 'not-started'}
            ]
        },
        {
            'name': '第四单元',
            'lessons': [
                {'number': '4.1', 'title': '分支限界法', 'status': 'in-progress'},
                {'number': '4.2', 'title': '优先队列应用', 'status': 'not-started'},
                {'number': '4.3', 'title': '最短路径问题', 'status': 'not-started'}
            ]
        },
        {
            'name': '第五单元',
            'lessons': [
                {'number': '5.1', 'title': '贪心法', 'status': 'completed'},
                {'number': '5.2', 'title': '贪心策略证明', 'status': 'not-started'},
                {'number': '5.3', 'title': '经典贪心问题', 'status': 'not-started'}
            ]
        },
        {
            'name': '第六单元',
            'lessons': [
                {'number': '6.1', 'title': '图论基础', 'status': 'not-started'},
                {'number': '6.2', 'title': '最小生成树', 'status': 'not-started'},
                {'number': '6.3', 'title': '网络流算法', 'status': 'not-started'}
            ]
        }
    ]

    # 统计任务点数量
    total_tasks_count = sum(len(unit['lessons']) for unit in course_units)
    completed_tasks_count = sum(
        1 for unit in course_units 
        for lesson in unit['lessons'] 
        if lesson['status'] == 'completed'
    )

    return render_template('client/learning_progress.html',
                         categories=categories,
                         has_progress=has_progress,
                         overall_completion=overall_completion,
                         completed_count=completed_count,
                         total_count=total_count,
                         difficult_count=difficult_count,
                         difficult_percentage=difficult_percentage,
                         study_hours=study_hours,
                         weekly_hours=weekly_hours,
                         daily_avg_hours=daily_avg_hours,
                         learning_path=learning_path,
                         difficult_contents=difficult_contents,
                         daily_study_data=daily_study_data,
                         mastery_distribution=mastery_distribution,
                         suggestions=suggestions,
                         course_units=course_units,
                         completed_tasks_count=completed_tasks_count,
                         total_tasks_count=total_tasks_count)


@client_bp.route('/rag-knowledge')
def rag_knowledge():
    """RAG知识库问答系统页面"""
    return render_template('client/rag_knowledge.html')

