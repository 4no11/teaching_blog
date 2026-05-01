"""
RAG知识库服务 - 支持多种LLM/Embedding提供商

支持的提供商：
1. Ollama (本地) - 免费，需要下载模型
2. OpenAI - 付费，API调用
3. Azure OpenAI - 企业版
4. HuggingFace - 免费，本地运行
5. 智谱AI (Zhipu) - 国内有免费额度
6. 自定义兼容接口

使用方式：
    通过环境变量或配置选择提供商
"""

import os
import json
import shutil
from datetime import datetime
from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


@dataclass
class KnowledgeBase:
    """知识库数据类"""
    id: str
    name: str
    description: str
    created_at: str
    document_count: int = 0
    path: str = ""
    chroma_path: str = ""
    user_id: int = None


class RAGKnowledgeService:
    """
    RAG知识库服务核心类 - 多模型支持版

    配置优先级（从高到低）：
    1. 环境变量
    2. config.json 配置文件
    3. 默认值（Ollama）
    """

    def __init__(self, provider: str = None):
        # 基础路径配置
        self.base_dir = os.path.join(os.getcwd(), 'rag_knowledge_base')
        self.kb_dir = os.path.join(self.base_dir, 'knowledge_bases')
        self.chroma_dir = os.path.join(self.base_dir, 'chroma_db')
        self.documents_dir = os.path.join(self.base_dir, 'documents')

        # 创建目录结构
        for directory in [self.base_dir, self.kb_dir, self.chroma_dir, self.documents_dir]:
            os.makedirs(directory, exist_ok=True)

        # 元数据存储文件
        self.metadata_file = os.path.join(self.base_dir, 'metadata.json')

        # 加载配置
        self.config = self._load_config()

        # 确定使用的提供商
        self.provider = provider or self.config.get('provider', 'ollama')
        logger.info(f"RAG服务将使用提供商: {self.provider}")

        # 初始化模型配置
        self._setup_provider_config()

        # 初始化嵌入模型和LLM
        self.embeddings = None
        self.llm = None

        # 加载元数据
        self._load_metadata()

        logger.info(f"RAG服务初始化完成，基础目录: {self.base_dir}")

    def _load_config(self) -> Dict:
        """加载配置文件"""
        config_file = os.path.join(self.base_dir, 'config.json')

        if os.path.exists(config_file):
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"配置文件读取失败: {e}")

        return {}

    def _save_config(self):
        """保存配置到文件"""
        config_file = os.path.join(self.base_dir, 'config.json')
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(self.config, f, ensure_ascii=False, indent=2)

    def _setup_provider_config(self):
        """根据提供商设置配置"""
        provider = self.provider.lower()

        if provider == 'ollama':
            self._setup_ollama()
        elif provider == 'openai':
            self._setup_openai()
        elif provider == 'azure':
            self._setup_azure_openai()
        elif provider == 'huggingface':
            self._setup_huggingface()
        elif provider == 'zhipu':
            self._setup_zhipu()
        else:
            logger.warning(f"未知提供商 '{provider}'，回退到 Ollama")
            self._setup_ollama()

    def _setup_ollama(self):
        """Ollama 配置 - 轻量级版本（适合8GB内存）"""
        self.ollama_base_url = os.environ.get(
            'OLLAMA_BASE_URL',
            self.config.get('ollama', {}).get('base_url', 'http://localhost:11434')
        )

        # 轻量级模型配置（总内存占用约3-4GB）
        self.llm_model = os.environ.get(
            'LLM_MODEL',
            self.config.get('ollama', {}).get('llm_model', 'qwen3:4b')
        )
        self.embedding_model = os.environ.get(
            'EMBEDDING_MODEL',
            self.config.get('ollama', {}).get('embedding_model', 'nomic-embed-text')
        )

        logger.info(f"[Ollama] LLM: {self.llm_model}, Embedding: {self.embedding_model}")

    def _setup_openai(self):
        """OpenAI API 配置"""
        self.openai_api_key = os.environ.get(
            'OPENAI_API_KEY',
            self.config.get('openai', {}).get('api_key', '')
        )
        self.openai_base_url = os.environ.get(
            'OPENAI_BASE_URL',
            self.config.get('openai', {}).get('base_url', None)
        )
        self.llm_model = os.environ.get(
            'LLM_MODEL',
            self.config.get('openai', {}).get('llm_model', 'gpt-4o-mini')
        )
        self.embedding_model = os.environ.get(
            'EMBEDDING_MODEL',
            self.config.get('openai', {}).get('embedding_model', 'text-embedding-3-small')
        )

        if not self.openai_api_key:
            raise ValueError("请设置 OPENAI_API_KEY 环境变量或在 config.json 中配置")

        logger.info(f"[OpenAI] LLM: {self.llm_model}, Embedding: {self.embedding_model}")

    def _setup_azure_openai(self):
        """Azure OpenAI 配置"""
        self.azure_api_key = os.environ.get(
            'AZURE_OPENAI_API_KEY',
            self.config.get('azure', {}).get('api_key', '')
        )
        self.azure_endpoint = os.environ.get(
            'AZURE_OPENAI_ENDPOINT',
            self.config.get('azure', {}).get('endpoint', '')
        )
        self.azure_deployment_name = os.environ.get(
            'AZURE_DEPLOYMENT_NAME',
            self.config.get('azure', {}).get('deployment_name', 'gpt-4')
        )
        self.azure_embedding_deployment = os.environ.get(
            'AZURE_EMBEDDING_DEPLOYMENT',
            self.config.get('azure', {}).get('embedding_deployment', 'text-embedding-ada-002')
        )
        self.azure_api_version = os.environ.get(
            'AZURE_API_VERSION',
            self.config.get('azure', {}).get('api_version', '2024-02-01')
        )

        if not self.azure_api_key or not self.azure_endpoint:
            raise ValueError("请设置 AZURE_OPENAI_API_KEY 和 AZURE_OPENAI_ENDPOINT")

        logger.info(f"[Azure] Endpoint: {self.azure_endpoint}, Deployment: {self.azure_deployment_name}")

    def _setup_huggingface(self):
        """HuggingFace 本地模型配置"""
        self.hf_llm_model = os.environ.get(
            'HF_LLM_MODEL',
            self.config.get('huggingface', {}).get('llm_model', 'THUDM/chatglm3-6b')
        )
        self.hf_embedding_model = os.environ.get(
            'HF_EMBEDDING_MODEL',
            self.config.get('huggingface', {}).get('embedding_model', 'BAAI/bge-large-zh-v1.5')
        )
        self.device = os.environ.get(
            'HF_DEVICE',
            self.config.get('huggingface', {}).get('device', 'cpu')
        )

        logger.info(f"[HuggingFace] LLM: {self.hf_llm_model}, Embedding: {self.hf_embedding_model}")

    def _setup_zhipu(self):
        """智谱AI配置（国产）"""
        self.zhipu_api_key = os.environ.get(
            'ZHIPUAI_API_KEY',
            self.config.get('zhipu', {}).get('api_key', '')
        )
        self.llm_model = os.environ.get(
            'LLM_MODEL',
            self.config.get('zhipu', {}).get('llm_model', 'glm-4-flash')
        )
        self.embedding_model = os.environ.get(
            'EMBEDDING_MODEL',
            self.config.get('zhipu', {}).get('embedding_model', 'embedding-3')
        )

        if not self.zhipu_api_key:
            raise ValueError("请设置 ZHIPUAI_API_KEY 环境变量或在 config.json 中配置")

        logger.info(f"[Zhipu] LLM: {self.llm_model}, Embedding: {self.embedding_model}")

    def _load_metadata(self):
        """加载知识库元数据"""
        if os.path.exists(self.metadata_file):
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                self.metadata = json.load(f)
        else:
            self.metadata = {'knowledge_bases': {}}
            self._save_metadata()

    def _save_metadata(self):
        """保存元数据"""
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(self.metadata, f, ensure_ascii=False, indent=2)

    def reload_metadata(self):
        """强制从磁盘重新加载元数据"""
        self._load_metadata()
        logger.info("📂 元数据已从磁盘重新加载")

    # ==================== 模型初始化 ====================

    def _init_embeddings(self):
        """延迟初始化嵌入模型（使用纯HTTP，避免LangChain Embeddings兼容性问题）"""
        if self.embeddings is not None:
            return

        provider = self.provider.lower()
        logger.info(f"正在初始化 [{provider}] 嵌入模型 (HTTP)...")

        if provider == 'openai':
            self.embeddings = 'openai_http'  # 标记，实际用 _get_embedding()

        elif provider == 'ollama':
            self.embeddings = 'ollama_http'  # 标记，实际用 _get_embedding()（避免langchain_ollama在Windows上的DLL崩溃问题）

        elif provider == 'azure':
            from langchain_openai import AzureOpenAIEmbeddings
            self.embeddings = AzureOpenAIEmbeddings(
                azure_deployment=self.azure_embedding_deployment,
                openai_api_version=self.azure_api_version,
                azure_endpoint=self.azure_endpoint,
                api_key=self.azure_api_key
            )

        elif provider == 'huggingface':
            from langchain_huggingface import HuggingFaceEmbeddings
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.hf_embedding_model,
                model_kwargs={'device': self.device}
            )

        elif provider == 'zhipu':
            from langchain_community.embeddings import ZhipuAIEmbeddings
            self.embeddings = ZhipuAIEmbeddings(
                zhipuai_api_key=self.zhipu_api_key,
                model=self.embedding_model
            )

        else:
            raise ValueError(f"不支持的提供商: {provider}")

        logger.info(f"✅ 嵌入模型初始化完成 ({provider})")

    def _get_embedding(self, text: str) -> list:
        """获取文本的向量嵌入（支持多种provider）"""
        import requests

        provider = self.provider.lower()

        if provider == 'ollama':
            # ✅ 使用 /api/embed 端点（/api/embeddings 在某些版本返回空数组）
            response = requests.post(
                f'{self.ollama_base_url}/api/embed',
                json={
                    'model': self.embedding_model,
                    'input': text
                },
                timeout=120  # 增加超时时间（本地模型可能较慢）
            )
            if response.status_code != 200:
                raise Exception(f'Ollama Embedding API错误: {response.text}')

            result = response.json()
            # /api/embed 返回格式: {"model":"...", "embeddings": [[...]]}
            embeddings = result.get('embeddings', [])
            if isinstance(embeddings, list) and len(embeddings) > 0:
                return embeddings[0]  # 返回第一个（也是唯一的）向量

            # 兼容旧格式: {"embedding": [...]}
            embedding = result.get('embedding', [])
            if embedding:
                return embedding

            raise Exception(f'Ollama Embedding API返回空向量: {result}')

        elif provider == 'openai':
            headers = {
                'Authorization': f'Bearer {self.openai_api_key}',
                'Content-Type': 'application/json'
            }
            base_url = self.openai_base_url or 'https://api.openai.com/v1'
            response = requests.post(
                f'{base_url}/embeddings',
                headers=headers,
                json={
                    'model': self.embedding_model,
                    'input': text
                },
                timeout=60
            )
            if response.status_code != 200:
                raise Exception(f'OpenAI Embedding API错误: {response.text}')
            result = response.json()
            return result['data'][0]['embedding']

        elif provider == 'azure':
            headers = {
                'api-key': self.azure_api_key,
                'Content-Type': 'application/json'
            }
            response = requests.post(
                f'{self.azure_endpoint}/openai/deployments/{self.azure_embedding_deployment}/embeddings?api-version={self.azure_api_version}',
                headers=headers,
                json={
                    'input': text
                },
                timeout=60
            )
            if response.status_code != 200:
                raise Exception(f'Azure Embedding API错误: {response.text}')
            result = response.json()
            return result['data'][0]['embedding']

        else:
            raise ValueError(f"不支持的 Embedding 提供商: {provider}")

    def _call_llm(self, messages: list) -> str:
        """调用 LLM 生成回答（支持多种provider，纯HTTP实现）"""
        import requests

        provider = self.provider.lower()

        if provider == 'ollama':
            response = requests.post(
                f'{self.ollama_base_url}/api/chat',
                json={
                    'model': self.llm_model,
                    'messages': messages,
                    'stream': False
                },
                timeout=120
            )
            if response.status_code != 200:
                raise Exception(f'Ollama Chat API错误: {response.text}')
            return response.json().get('message', {}).get('content', '')

        elif provider == 'openai':
            headers = {
                'Authorization': f'Bearer {self.openai_api_key}',
                'Content-Type': 'application/json'
            }
            base_url = self.openai_base_url or 'https://api.openai.com/v1'
            response = requests.post(
                f'{base_url}/chat/completions',
                headers=headers,
                json={
                    'model': self.llm_model,
                    'messages': messages,
                    'temperature': 0.7,
                    'max_tokens': 2048
                },
                timeout=120
            )
            if response.status_code != 200:
                raise Exception(f'OpenAI Chat API错误: {response.text}')
            result = response.json()
            return result['choices'][0]['message']['content']

        elif provider == 'azure':
            headers = {
                'api-key': self.azure_api_key,
                'Content-Type': 'application/json'
            }
            response = requests.post(
                f'{self.azure_endpoint}/openai/deployments/{self.azure_deployment_name}/chat/completions?api-version={self.azure_api_version}',
                headers=headers,
                json={
                    'messages': messages,
                    'temperature': 0.7,
                    'max_tokens': 2048
                },
                timeout=120
            )
            if response.status_code != 200:
                raise Exception(f'Azure Chat API错误: {response.text}')
            result = response.json()
            return result['choices'][0]['message']['content']

        else:
            raise ValueError(f"不支持的 LLM 提供商: {provider}")

    def _init_llm(self):
        """延迟初始化LLM（使用纯HTTP调用，基于LangChain设计模式）"""
        if self.llm is not None:
            return

        provider = self.provider.lower()
        logger.info(f"正在初始化 [{provider}] LLM (HTTP + LangChain Pattern)...")

        if provider == 'openai':
            self.llm = 'openai_http'  # 标记，实际用 _call_llm()

        elif provider == 'ollama':
            self.llm = 'ollama_http'  # 标记，实际用 _call_llm()（避免langchain_ollama在Windows上的DLL崩溃问题）

        elif provider == 'azure':
            from langchain_openai import AzureChatOpenAI
            self.llm = AzureChatOpenAI(
                azure_deployment=self.azure_deployment_name,
                openai_api_version=self.azure_api_version,
                azure_endpoint=self.azure_endpoint,
                api_key=self.azure_api_key,
                temperature=0.7
            )

        elif provider == 'huggingface':
            from langchain_huggingface import HuggingFacePipeline
            from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

            logger.warning("[HuggingFace] 加载本地大型语言模型可能需要较长时间和大量内存...")

            tokenizer = AutoTokenizer.from_pretrained(self.hf_llm_model)
            model = AutoModelForCausalLM.from_pretrained(self.hf_llm_model)

            pipe = pipeline(
                "text-generation",
                model=model,
                tokenizer=tokenizer,
                max_new_tokens=1024,
                temperature=0.7,
                device=self.device
            )
            self.llm = HuggingFacePipeline(pipeline=pipe)

        elif provider == 'zhipu':
            from langchain_community.chat_models import ChatZhipuAI
            self.llm = ChatZhipuAI(
                zhipuai_api_key=self.zhipu_api_key,
                model=self.llm_model,
                temperature=0.7
            )

        else:
            raise ValueError(f"不支持的提供商: {provider}")

        logger.info(f"✅ LLM 初始化完成 ({provider})")

    # ==================== 知识库管理 ====================

    def create_knowledge_base(self, name: str, description: str = "", user_id: int = None) -> Dict:
        """创建新的知识库"""
        self.reload_metadata()

        # 检查同一用户下是否有重名
        for kb_id, kb_info in self.metadata['knowledge_bases'].items():
            if kb_info['name'] == name and kb_info.get('user_id') == user_id:
                logger.warning(f"知识库名称已存在: {name} (ID: {kb_id})")
                return {
                    'success': False,
                    'error': f'知识库 "{name}" 已存在，请使用其他名称',
                    'existing_kb': KnowledgeBase(**kb_info).__dict__
                }

        kb_id = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + name.replace(' ', '_')

        kb_path = os.path.join(self.kb_dir, kb_id)
        os.makedirs(kb_path, exist_ok=True)

        chroma_path = os.path.join(self.chroma_dir, kb_id)

        kb_info = {
            'id': kb_id,
            'name': name,
            'description': description,
            'created_at': datetime.now().isoformat(),
            'document_count': 0,
            'path': kb_path,
            'chroma_path': chroma_path,
            'user_id': user_id
        }

        self.metadata['knowledge_bases'][kb_id] = kb_info
        self._save_metadata()

        logger.info(f"✅ 知识库创建成功: {name} (ID: {kb_id}, User: {user_id})")

        return {
            'success': True,
            'knowledge_base': KnowledgeBase(**kb_info).__dict__
        }

    def delete_knowledge_base(self, kb_id: str) -> Dict:
        """删除知识库及其所有数据"""
        if kb_id not in self.metadata['knowledge_bases']:
            return {'success': False, 'error': '知识库不存在'}

        kb_info = self.metadata['knowledge_bases'][kb_id]

        try:
            if os.path.exists(kb_info['path']):
                shutil.rmtree(kb_info['path'])

            if os.path.exists(kb_info.get('chroma_path', '')):
                shutil.rmtree(kb_info['chroma_path'])

            del self.metadata['knowledge_bases'][kb_id]
            self._save_metadata()

            logger.info(f"✅ 知识库删除成功: {kb_id}")
            return {'success': True}

        except Exception as e:
            logger.error(f"❌ 删除知识库失败: {e}")
            return {'success': False, 'error': str(e)}

    def list_knowledge_bases(self, user_id: int = None) -> List[Dict]:
        """获取知识库列表（支持按用户过滤）"""
        self.reload_metadata()

        kbs = []
        for kb_id, info in self.metadata['knowledge_bases'].items():
            # 如果指定了user_id，只返回该用户的
            if user_id is not None and info.get('user_id') != user_id:
                continue
            kbs.append(KnowledgeBase(**info).__dict__)
        return kbs

    def get_knowledge_base(self, kb_id: str) -> Optional[Dict]:
        """获取单个知识库详情"""
        if kb_id in self.metadata['knowledge_bases']:
            return KnowledgeBase(**self.metadata['knowledge_bases'][kb_id]).__dict__
        return None

    # ==================== 文档上传与处理 ====================

    def upload_documents(
        self,
        kb_id: str,
        files: List[Tuple[str, bytes]],
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        skip_vectorization: bool = False
    ) -> Dict:
        """上传文档（支持快速模式和完整模式）"""
        if kb_id not in self.metadata['knowledge_bases']:
            return {'success': False, 'error': '知识库不存在'}

        kb_info = self.metadata['knowledge_bases'][kb_id]
        doc_dir = os.path.join(self.documents_dir, kb_id)
        os.makedirs(doc_dir, exist_ok=True)

        uploaded_files = []
        all_documents = []

        try:
            for file_item in files:
                if len(file_item) == 3:
                    filename, content, display_name = file_item
                else:
                    filename, content = file_item
                    display_name = filename
                filepath = os.path.join(doc_dir, filename)
                with open(filepath, 'wb') as f:
                    f.write(content)

                uploaded_files.append({
                    'filename': display_name,
                    'path': filepath,
                    'size': len(content)
                })

                documents = self._load_document(filepath, filename)
                if documents:
                    all_documents.extend(documents)

            if not all_documents:
                return {
                    'success': False,
                    'error': '无法解析任何文档',
                    'uploaded_files': uploaded_files
                }

            self.metadata['knowledge_bases'][kb_id]['document_count'] += len(uploaded_files)
            self._save_metadata()

            if skip_vectorization:
                logger.info(f"✅ 文件已保存（跳过向量化）: {len(uploaded_files)} 个文件")
                return {
                    'success': True,
                    'uploaded_files': uploaded_files,
                    'total_chunks': 0,
                    'message': f'成功上传{len(uploaded_files)}个文件（已保存，待向量化）',
                    'vectorized': False
                }

            logger.info(f"开始向量化处理 {len(all_documents)} 个文档...")

            # ========== 初始化进度追踪 ==========
            from services.vectorization_progress import get_progress_tracker
            progress_tracker = get_progress_tracker()
            progress_tracker.start_task(kb_id, total_chunks=0)  # 先初始化，后面更新总数

            # ========== 文本切片（使用纯Python实现，避免DLL问题）==========
            progress_tracker.update_progress(kb_id, message="正在切割文本...")

            from services.pure_splitter import PurePythonTextSplitter

            text_splitter = PurePythonTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )

            chunks = text_splitter.split_documents(all_documents)
            logger.info(f"文档已切分为 {len(chunks)} 个文本块")

            if not chunks:
                progress_tracker.fail_task(kb_id, "文本切片后无有效内容")
                return {
                    'success': False,
                    'error': '文本切片后无有效内容',
                    'uploaded_files': uploaded_files
                }

            # 更新总块数
            progress_tracker.update_progress(kb_id, current=0, total_chunks=len(chunks), message=f"准备生成 {len(chunks)} 个向量...")

            # ========== 使用 Embeddings + ChromaDB兼容层 ==========
            from services.chroma_compat import PersistentClient as ChromaPersistentClient

            # 初始化 Embeddings（OpenAI 用 HTTP，其他用 LangChain）
            self._init_embeddings()

            chroma_path = kb_info.get('chroma_path',
                                      os.path.join(self.chroma_dir, kb_id))
            os.makedirs(chroma_path, exist_ok=True)

            logger.info(f"正在生成 {len(chunks)} 个向量并存储到ChromaDB...")

            # 创建或获取Chroma集合（使用兼容层）
            client = ChromaPersistentClient(path=chroma_path)
            collection = client.get_or_create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"},
                embedding_function=None  # 关键：禁用默认embedding，手动提供向量
            )

            # 批量生成向量并存储
            batch_size = 10
            successful_chunks = 0
            failed_chunks = 0
            total_chunks = len(chunks)

            logger.info(f"开始生成 {total_chunks} 个向量...")
            
            use_langchain_embedding = self.embeddings not in ('openai_http', 'ollama_http')

            for batch_start in range(0, total_chunks, batch_size):
                batch_end = min(batch_start + batch_size, total_chunks)
                batch_chunks = chunks[batch_start:batch_end]
                batch_ids = [f"{kb_id}_{i}" for i in range(batch_start, batch_end)]

                try:
                    # 生成向量（根据类型选择方法）
                    if use_langchain_embedding:
                        batch_texts = [chunk.page_content for chunk in batch_chunks]
                        batch_embeddings = self.embeddings.embed_documents(batch_texts)
                    else:
                        batch_embeddings = []
                        for chunk in batch_chunks:
                            embedding = self._get_embedding(chunk.page_content)
                            if embedding:
                                batch_embeddings.append(embedding)
                            else:
                                batch_embeddings.append(None)
                    for i, (chunk, embedding) in enumerate(zip(batch_chunks, batch_embeddings)):
                        chunk_id = batch_ids[i]
                        if embedding and len(embedding) > 0:
                            collection.upsert(
                                ids=[chunk_id],
                                documents=[chunk.page_content],
                                metadatas=[{
                                    **chunk.metadata,
                                    'kb_id': kb_id,
                                    'chunk_index': batch_start + i,
                                    'source': chunk.metadata.get('source', '')
                                }],
                                embeddings=[embedding]
                            )
                            successful_chunks += 1
                        else:
                            failed_chunks += 1
                            logger.warning(f"Chunk {batch_start+i}: 返回空向量")

                    progress_tracker.update_progress(
                        kb_id,
                        current=batch_end,
                        total_chunks=total_chunks,
                        message=f"正在生成向量 {batch_end}/{total_chunks} (成功:{successful_chunks}, 失败:{failed_chunks})..."
                    )

                except Exception as e:
                    logger.error(f"批次 {batch_start}-{batch_end} 向量化失败: {e}")
                    failed_chunks += len(batch_chunks)
                    
                    import time
                    time.sleep(2)

            # 向量化完成（无论是否有失败）
            final_message = f"向量化完成! 成功: {successful_chunks}/{total_chunks}"
            if failed_chunks > 0:
                final_message += f", 跳过: {failed_chunks}个失败块"

            logger.info(final_message)

            # 标记任务完成（即使部分失败也算完成）
            progress_tracker.complete_task(kb_id, result={
                'total_chunks': total_chunks,
                'successful_chunks': successful_chunks,
                'failed_chunks': failed_chunks
            })

            return {
                'success': True,
                'uploaded_files': uploaded_files,
                'total_chunks': total_chunks,
                'successful_chunks': successful_chunks,
                'failed_chunks': failed_chunks,
                'message': f'成功上传{len(uploaded_files)}个文件，向量化{successful_chunks}/{total_chunks}个文本块',
                'vectorized': True
            }

        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            import traceback
            traceback.print_exc()

            # 标记任务失败
            try:
                progress_tracker.fail_task(kb_id, str(e))
            except:
                pass

            return {'success': False, 'error': str(e)}

    def _load_document(self, filepath: str, filename: str) -> List:
        """根据文件类型加载文档 - 纯Python版本（避免PyTorch DLL问题）"""
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        logger.info(f"📄 正在加载文档: {filename} (类型: {ext})")

        try:
            # 检查文件
            if not os.path.exists(filepath):
                logger.error(f"❌ 文件不存在: {filepath}")
                return []

            file_size = os.path.getsize(filepath)
            if file_size == 0:
                logger.warning(f"⚠️ 文件为空: {filename}")
                return []

            documents = []

            # ========== 纯Python读取文本文件（完全绕过PyTorch）==========
            if ext in ['txt', 'md', 'csv', 'log', 'json', 'xml', 'html']:
                content = None
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        with open(filepath, 'r', encoding=encoding) as f:
                            content = f.read()
                        logger.info(f"✅ 纯Python读取成功 [{encoding}]")
                        break
                    except (UnicodeDecodeError, Exception):
                        continue

                if content and len(content.strip()) > 0:
                    from services.pure_splitter import Document
                    documents.append(Document(
                        page_content=content,
                        metadata={'source': filename}
                    ))
                    logger.info(f"文本加载成功: {len(content)} 字符")
                else:
                    logger.error(f"❌ 文件内容为空或无法解码")
                    return []

            # ========== Word文档（使用python-docx）==========
            elif ext in ['docx', 'doc']:
                try:
                    import docx
                    doc = docx.Document(filepath)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    if paragraphs:
                        content = '\n'.join(paragraphs)
                        from services.pure_splitter import Document
                        documents.append(Document(
                            page_content=content,
                            metadata={'source': filename}
                        ))
                        logger.info(f"Word加载成功: {len(content)} 字符")
                    else:
                        logger.warning("⚠️ Word文档内容为空")
                        return []
                except ImportError:
                    logger.warning("python-docx未安装")
                    return []
                except Exception as e:
                    logger.error(f"❌ Word加载失败: {e}")
                    return []

            # ========== PDF（仅当需要时才导入PyPDF）==========
            elif ext == 'pdf':
                try:
                    import pypdf
                    reader = pypdf.PdfReader(filepath)
                    text_parts = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(text.strip())
                    if text_parts:
                        content = '\n\n'.join(text_parts)
                        from services.pure_splitter import Document
                        documents.append(Document(
                            page_content=content,
                            metadata={'source': filename}
                        ))
                        logger.info(f"✅ PDF加载成功: {len(reader.pages)} 页, {len(content)} 字符")
                    else:
                        logger.warning("⚠️ PDF文档内容为空或无法提取文字（可能是扫描件）")
                        return []
                except ImportError:
                    logger.warning("pypdf未安装，请运行: pip install pypdf")
                    return []
                except Exception as e:
                    logger.error(f"❌ PDF加载失败: {e}")
                    return []

            # ========== 其他格式：最后尝试 ==========
            else:
                try:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    if content.strip():
                        from services.pure_splitter import Document
                        documents.append(Document(
                            page_content=content,
                            metadata={'source': filename}
                        ))
                        logger.info("兜底读取成功")
                except Exception as e:
                    logger.error(f"❌ 所有方法失败: {e}")
                    return []

            # 添加元数据
            for doc in documents:
                if 'source' not in doc.metadata:
                    doc.metadata['source'] = filename

            logger.info(f"✅ 最终成功: {len(documents)} 个文档段")
            return documents

        except Exception as e:
            logger.error(f"❌ 加载失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return []

    def get_documents(self, kb_id: str) -> List[Dict]:
        """获取知识库中的文档列表"""
        doc_dir = os.path.join(self.documents_dir, kb_id)

        if not os.path.exists(doc_dir):
            return []

        documents = []
        for filename in os.listdir(doc_dir):
            filepath = os.path.join(doc_dir, filename)
            if os.path.isfile(filepath):
                stat = os.stat(filepath)
                documents.append({
                    'filename': filename,
                    'size': stat.st_size,
                    'modified_time': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                    'size_formatted': self._format_size(stat.st_size)
                })

        return sorted(documents, key=lambda x: x['modified_time'], reverse=True)

    def delete_document(self, kb_id: str, filename: str) -> Dict:
        """删除指定文档"""
        doc_dir = os.path.join(self.documents_dir, kb_id)
        filepath = os.path.join(doc_dir, filename)

        if not os.path.exists(filepath):
            return {'success': False, 'error': '文件不存在'}

        try:
            os.remove(filepath)

            if kb_id in self.metadata['knowledge_bases']:
                current_count = self.metadata['knowledge_bases'][kb_id].get('document_count', 0)
                self.metadata['knowledge_bases'][kb_id]['document_count'] = max(0, current_count - 1)
                self._save_metadata()

            return {'success': True}

        except Exception as e:
            return {'success': False, 'error': str(e)}

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"

    # ==================== RAG问答 ====================

    def chat(
        self,
        kb_id: str,
        question: str,
        top_k: int = 5,
        return_sources: bool = True
    ) -> Dict:
        """RAG问答 - 带引用来源的智能回答（完全使用ChromaDB原生API）"""
        if kb_id not in self.metadata['knowledge_bases']:
            return {'success': False, 'error': '知识库不存在'}

        try:
            from services.chroma_compat import PersistentClient as ChromaPersistentClient
            import requests

            chroma_path = self.metadata['knowledge_bases'][kb_id].get(
                'chroma_path',
                os.path.join(self.chroma_dir, kb_id)
            )

            if not os.path.exists(chroma_path):
                return {
                    'success': False,
                    'error': '该知识库暂无文档数据，请先上传文档并完成向量化'
                }

            # 使用ChromaDB兼容层查询
            client = ChromaPersistentClient(path=chroma_path)
            collection = client.get_collection("documents")

            # 生成问题向量（OpenAI 用 HTTP，其他用 LangChain）
            self._init_embeddings()  # 确保已初始化
            if self.embeddings in ('openai_http', 'ollama_http'):
                question_embedding = self._get_embedding(question)
            else:
                question_embedding = self.embeddings.embed_query(question)

            if not question_embedding:
                return {'success': False, 'error': '无法生成向量'}

            # ========== 混合检索：向量搜索 + 关键词匹配 ==========
            # 原因：nomic-embed-text对中文语义检索能力有限，纯向量搜索可能漏掉相关文档

            total_count = collection.count()

            # 小数据集策略：全量获取后混合排序；大数据集：取top N
            if total_count <= 100:
                fetch_k = total_count
            else:
                fetch_k = min(max(top_k * 5, 30), total_count)

            vector_results = collection.query(
                query_embeddings=[question_embedding],
                n_results=fetch_k,
                include=['documents', 'metadatas', 'distances']
            )

            if not vector_results['documents'] or not vector_results['documents'][0]:
                return {
                    'success': True,
                    'answer': '抱歉，我在知识库中没有找到与您问题相关的内容。',
                    'sources': []
                }

            # 提取问题关键词（去除停用词和常见疑问词）
            import re
            stop_words = {'的', '是', '什么', '怎么', '如何', '哪', '这', '那', '有', '在', '和',
                          '与', '或', '了', '吗', '呢', '啊', '吧', '哦', '哈', '嗯', '详细', '请',
                          '告诉我', '介绍', '说明', '描述', '列出', '包含', '关于'}
            raw_chars = set(re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', question))
            keywords = raw_chars - stop_words

            # 扩展关键词：拆分常见复合词（如"实验步骤"->["实验","步骤"]）
            expanded_keywords = list(keywords)
            for kw in list(keywords):
                if len(kw) >= 4:
                    for i in range(len(kw) - 1):
                        expanded_keywords.append(kw[i:i+2])
            keywords = [kw for kw in expanded_keywords if len(kw) >= 1]

            # 混合评分：向量相似度 + 关键词匹配度
            scored_docs = []
            for i, (doc, metadata, distance) in enumerate(zip(
                vector_results['documents'][0],
                vector_results['metadatas'][0],
                vector_results['distances'][0]
            )):
                doc_text = doc.strip()
                if len(doc_text) < 30:
                    continue

                vector_score = 1 - distance

                keyword_score = 0
                if keywords:
                    matched = sum(1 for kw in keywords if kw in doc_text)
                    keyword_score = matched / len(keywords) if keywords else 0

                combined_score = 0.3 * vector_score + 0.7 * keyword_score

                scored_docs.append({
                    'doc': doc_text,
                    'metadata': metadata,
                    'vector_score': vector_score,
                    'keyword_score': keyword_score,
                    'combined_score': combined_score,
                    'distance': distance
                })

            # 按综合评分排序，取top_k
            scored_docs.sort(key=lambda x: x['combined_score'], reverse=True)
            scored_docs = scored_docs[:top_k]

            if not scored_docs:
                return {
                    'success': True,
                    'answer': '抱歉，我在知识库中没有找到与您问题相关的内容。',
                    'sources': []
                }

            # 构建上下文
            context_parts = []
            sources = []

            for idx, item in enumerate(scored_docs):
                doc = item['doc']
                meta = item['metadata']
                source_name = meta.get('source', f'文档{idx+1}')
                similarity_score = round(item['vector_score'] * 100, 1)

                context_parts.append(f"[文档{idx+1}] ({source_name})\n{doc}\n")

                sources.append({
                    'source': source_name,
                    'content': doc[:200] + ('...' if len(doc) > 200 else ''),
                    'score': similarity_score,
                    'page': meta.get('page', None)
                })

            context = "\n".join(context_parts)

            # 使用 LangChain 设计模式进行回答（OpenAI 用 HTTP，其他用 LangChain）
            self._init_llm()

            if self.llm in ('openai_http', 'ollama_http'):
                system_prompt = """你是一个专业的教育知识助手。请仔细阅读以下参考资料，基于资料内容详细回答用户的问题。

【重要原则】
1. 参考资料中包含的信息是回答的唯一依据，但你要主动提取和整合信息
2. 如果参考资料中确实没有相关信息，才说明找不到
3. 回答要详细、完整、有条理，尽量覆盖资料中的所有相关信息
4. 使用中文回答"""

                human_prompt = """【参考资料】
{context}

【用户问题】
{question}

【你的回答】"""

                prompt = human_prompt.format(context=context, question=question)
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ]
                
                result = self._call_llm(messages)
            
            else:
                # 使用 LangChain Chain（原生支持）
                from langchain_core.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
                from langchain_core.output_parsers import StrOutputParser

                system_prompt = """你是一个专业的教育知识助手。请仔细阅读以下参考资料，基于资料内容详细回答用户的问题。

【重要原则】
1. 参考资料中包含的信息是回答的唯一依据，但你要主动提取和整合信息
2. 如果参考资料中确实没有相关信息，才说明找不到
3. 回答要详细、完整、有条理，尽量覆盖资料中的所有相关信息
4. 使用中文回答"""

                human_prompt = """【参考资料】
{context}

【用户问题】
{question}

【你的回答】"""

                prompt = ChatPromptTemplate.from_messages([
                    SystemMessagePromptTemplate.from_template(system_prompt),
                    HumanMessagePromptTemplate.from_template(human_prompt)
                ])

                chain = prompt | self.llm | StrOutputParser()
                result = chain.invoke({
                    "context": context,
                    "question": question
                })

            avg_similarity = None
            valid_scores = [s['score'] for s in sources if s['score'] is not None]
            if valid_scores:
                avg_similarity = round(sum(valid_scores) / len(valid_scores), 1)

            response = {
                'success': True,
                'answer': result,
                'question': question,
                'sources': sources if return_sources else [],
                'avg_similarity': avg_similarity,
                'model_used': f"{self.provider}:{self.llm_model}",
                'timestamp': datetime.now().isoformat(),
                'retrieved_docs': len(sources)
            }

            logger.info(f"✅ 问答完成，检索到 {len(sources)} 个相关文档")
            return response

        except Exception as e:
            logger.error(f"❌ 问答过程出错: {e}")
            import traceback
            traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def chat_stream(self, kb_id: str, question: str, top_k: int = 5):
        """流式问答（完全使用ChromaDB原生API）"""
        if kb_id not in self.metadata['knowledge_bases']:
            yield json.dumps({'error': '知识库不存在'}, ensure_ascii=False)
            return

        try:
            import chromadb
            import requests

            chroma_path = self.metadata['knowledge_bases'][kb_id].get(
                'chroma_path',
                os.path.join(self.chroma_dir, kb_id)
            )

            if not os.path.exists(chroma_path):
                yield json.dumps({'error': '该知识库暂无文档数据'}, ensure_ascii=False)
                return

            # 使用ChromaDB兼容层查询
            client = ChromaPersistentClient(path=chroma_path)
            collection = client.get_collection("documents")

            # 生成问题的向量（支持多种provider）
            try:
                question_embedding = self._get_embedding(question)
                if not question_embedding:
                    yield json.dumps({'error': '无法生成问题向量'}, ensure_ascii=False)
                    return
            except Exception as e:
                yield json.dumps({'error': f'Embedding API错误: {str(e)}'}, ensure_ascii=False)
                return

            # 在ChromaDB中查询相似文档
            results = collection.query(
                query_embeddings=[question_embedding],
                n_results=top_k,
                include=['documents', 'metadatas', 'distances']
            )

            context_parts = []
            sources = []

            for i, (doc, metadata, distance) in enumerate(zip(
                results['documents'][0] if results['documents'] else [],
                results['metadatas'][0] if results['metadatas'] else [],
                results['distances'][0] if results['distances'] else []
            )):
                source_name = metadata.get('source', f'文档{i+1}')
                similarity_score = round((1 - distance) * 100, 1)
                context_parts.append(f"[文档{i+1}] ({source_name})\n{doc}\n")
                sources.append({
                    'source': source_name,
                    'content': doc[:200],
                    'score': similarity_score,
                    'page': metadata.get('page', None)
                })

            context = "\n".join(context_parts)

            prompt = f"""你是一个专业的教育知识助手。请根据以下参考资料回答问题。

【参考资料】
{context}

【问题】
{question}

【回答】"""

            # 初始化LLM进行流式输出
            self._init_llm()

            if self.llm in ('openai_http', 'ollama_http'):
                import requests
                provider = self.provider.lower()
                if provider == 'ollama':
                    response = requests.post(
                        f'{self.ollama_base_url}/api/chat',
                        json={
                            'model': self.llm_model,
                            'messages': [{'role': 'user', 'content': prompt}],
                            'stream': True
                        },
                        timeout=120,
                        stream=True
                    )
                    for line in response.iter_lines():
                        if line:
                            data = json.loads(line)
                            content = data.get('message', {}).get('content', '')
                            if content:
                                yield json.dumps({
                                    'type': 'content',
                                    'content': content,
                                    'sources': sources
                                }, ensure_ascii=False)
                else:
                    headers = {
                        'Authorization': f'Bearer {self.openai_api_key}',
                        'Content-Type': 'application/json'
                    }
                    base_url = self.openai_base_url or 'https://api.openai.com/v1'
                    response = requests.post(
                        f'{base_url}/chat/completions',
                        headers=headers,
                        json={
                            'model': self.llm_model,
                            'messages': [{'role': 'user', 'content': prompt}],
                            'stream': True
                        },
                        timeout=120,
                        stream=True
                    )
                    for line in response.iter_lines():
                        if line and line.startswith('data: '):
                            data_str = line[6:]
                            if data_str.strip() == '[DONE]':
                                break
                            data = json.loads(data_str)
                            content = data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield json.dumps({
                                    'type': 'content',
                                    'content': content,
                                    'sources': sources
                                }, ensure_ascii=False)
            else:
                for chunk in self.llm.stream(prompt):
                    yield json.dumps({
                        'type': 'content',
                        'content': chunk.content,
                        'sources': sources
                    }, ensure_ascii=False)

            yield json.dumps({'type': 'done'}, ensure_ascii=False)

        except Exception as e:
            yield json.dumps({'type': 'error', 'error': str(e)}, ensure_ascii=False)


# 全局单例实例
rag_service = None


def get_rag_service(provider: str = None) -> RAGKnowledgeService:
    """获取RAG服务实例（支持指定提供商）"""
    global rag_service
    if rag_service is None:
        rag_service = RAGKnowledgeService(provider=provider)
    return rag_service
