"""
快速测试：不依赖PyTorch的文档解析
"""

import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

print("=" * 60)
print("   文档解析测试（无PyTorch依赖版）")
print("=" * 60)

# 测试文件
docx_file = r'd:\boke\teaching_blog\rag_knowledge_base\documents\20260501172221_爬虫\2..docx'
txt_file = r'd:\boke\teaching_blog\rag_knowledge_base\documents\20260501172221_爬虫\test_document.txt'

# ====== 测试1: TXT文件 ======
print("\n[测试1] TXT文件...")
if os.path.exists(txt_file):
    try:
        with open(txt_file, 'r', encoding='utf-8') as f:
            content = f.read()
        print(f"✅ TXT读取成功！")
        print(f"   长度: {len(content)} 字符")
        print(f"   前100字: {content[:100]}")
    except Exception as e:
        print(f"❌ 失败: {e}")
else:
    print("⚠️  文件不存在")

# ====== 测试2: DOCX文件 ======
print("\n[测试2] DOCX文件 (使用python-docx)...")
if os.path.exists(docx_file):
    try:
        import docx
        doc = docx.Document(docx_file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n'.join(paragraphs)
        print(f"✅ DOCX读取成功！")
        print(f"   段落数: {len(paragraphs)}")
        print(f"   总长度: {len(content)} 字符")
        if content:
            print(f"   前150字: {content[:150]}")
    except ImportError:
        print("❌ python-docx 未安装")
        print("   安装: pip install python-docx")
    except Exception as e:
        print(f"❌ 失败: {e}")
else:
    print("⚠️  文件不存在")

# ====== 测试3: 创建Document对象 ======
print("\n[测试3] 创建LangChain Document对象...")

try:
    from langchain.schema import Document
    test_doc = Document(
        page_content="这是测试内容",
        metadata={'source': 'test.txt'}
    )
    print(f"✅ Document对象创建成功!")
    print(f"   page_content: {test_doc.page_content}")
    print(f"   metadata: {test_doc.metadata}")
except Exception as e:
    print(f"❌ 失败: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "=" * 60)
print("测试完成!")
